from flask import Flask, render_template, request, send_file
from io import BytesIO
from docx import Document
from docx.shared import Pt,Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

app = Flask(__name__,template_folder='template')

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/submit', methods=['POST'])
    
def submit():
    release_version = request.form['release_version']
    reference_ticket = request.form['reference_ticket']
    release_description = request.form['release_description']
    type_of_change = request.form['type_of_change']
    project_name = request.form['project_name']
    change_areas = request.form.getlist('change_areas')
    deployment_steps = request.form['deployment_steps']
    
    # Get the desired filename from the form
    desired_filename = request.form['filename']
    
    # Create a BytesIO object to store the document in memory
    output = BytesIO()

    # Create a new Word document
    document = Document()

    # Set the document font size
    document.styles['Normal'].font.size = Pt(12)

    # Add the centered heading with a single line gap
    heading = document.add_heading('Release Note', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph()

    # Release Version
    release_version_paragraph = document.add_paragraph()
    release_version_paragraph.add_run('Release Version: ').bold = True
    release_version_paragraph.add_run(release_version)
    
    # Release Version
    reference_ticket_paragraph = document.add_paragraph()
    reference_ticket_paragraph.add_run('Reference Ticket: ').bold = True
    reference_ticket_paragraph.add_run(reference_ticket)	

    # Release Description
    release_description_paragraph = document.add_paragraph()
    release_description_paragraph.add_run('Release Description: ').bold = True
    release_description_paragraph.add_run(release_description)

    # Type of Change
    type_of_change_paragraph = document.add_paragraph()
    type_of_change_paragraph.add_run('Type of Change: ').bold = True
    type_of_change_paragraph.add_run(type_of_change)

    # Project Name
    project_name_paragraph = document.add_paragraph()
    project_name_paragraph.add_run('Project Name: ').bold = True
    project_name_paragraph.add_run(project_name)

    # Change Areas
    change_areas_paragraph = document.add_paragraph()
    change_areas_paragraph.add_run('Change Areas:').bold = True
    change_areas_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    change_areas_paragraph.paragraph_format.space_after = Pt(6)

    # Add the selected change areas as checked checkboxes
    for area in ['project_changes', 'database_changes', 'configuration_changes']:
        checkbox_paragraph = document.add_paragraph()
        checkbox = checkbox_paragraph.add_run('\t\u2611' if area in change_areas else '\t\u2610')
        checkbox.bold = True
        checkbox_paragraph.add_run(f" {area.replace('_', ' ').title()}")

    # Gap before Deployment Steps
    document.add_paragraph()

    # Deployment Steps
    deployment_steps_paragraph = document.add_paragraph()
    deployment_steps_paragraph.add_run('Deployment Steps:').bold = True
    deployment_steps_paragraph.paragraph_format.space_after = Pt(6)

    # Add a new paragraph to start the deployment steps on the next line
    deployment_steps_paragraph.add_run().add_break()
    
    # Add the deployment steps preserving formatting and indentation
    steps = re.split(r'\r?\n', deployment_steps)
    for step in steps:
        if step.strip():
            # Add each step as a separate run within the same paragraph
            run_step = deployment_steps_paragraph.add_run(step + '\n').font.size = Pt(11)
    
    # Save the Word document to the BytesIO object
    document.save(output)

    # Seek to the beginning of the BytesIO stream
    output.seek(0)
    
    # Construct the filename using the desired filename input from the form
    filename = desired_filename.strip() + ".docx" if desired_filename.strip() else "ReleaseNote.docx"
    
    return send_file(output, as_attachment=True,
                     download_name=filename,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                     etag=False,
                     conditional=False)
if __name__ == '__main__':
    app.run(debug=True)
    
    